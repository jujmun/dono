from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "dono-ico-childrens-code-assessment-2026-08-06.md"
OUTPUT = ROOT / "dono-ico-childrens-code-assessment-2026-08-06.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "666666"
BLACK = "000000"


def set_font(run, name="Calibri", size=11, color=BLACK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style(style, size, color, bold, before, after, line=1.10):
    style.font.name = "Calibri"
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    pf = style.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.keep_with_next = True if style.name.startswith("Heading") else None


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "22")
    rpr.extend([fonts, color, underline, size])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([rpr, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_inline(paragraph, text):
    token = re.compile(r"(\*\*.+?\*\*|`.+?`|https?://\S+)")
    pos = 0
    for match in token.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_font(run)
        item = match.group(0)
        if item.startswith("**"):
            run = paragraph.add_run(item[2:-2])
            set_font(run, bold=True)
        elif item.startswith("`"):
            run = paragraph.add_run(item[1:-1])
            set_font(run, name="Menlo", size=9.5, color="333333")
        else:
            url = item.rstrip(".,;)")
            tail = item[len(url):]
            add_hyperlink(paragraph, "Official source", url)
            if tail:
                run = paragraph.add_run(tail)
                set_font(run)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_font(run)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_text, fld_end])


def add_numbering_definition(doc, fmt, text, num_id, abstract_id):
    numbering = doc.part.numbering_part.element
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    numfmt = OxmlElement("w:numFmt")
    numfmt.set(qn("w:val"), fmt)
    lvltext = OxmlElement("w:lvlText")
    lvltext.set(qn("w:val"), text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.extend([tabs, ind, spacing])
    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    rpr.append(fonts)
    lvl.extend([start, numfmt, lvltext, suff, ppr, rpr])
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)


def apply_num(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    numpr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    numpr.extend([ilvl, numid])
    ppr.append(numpr)


def main():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    set_style(doc.styles["Normal"], 11, BLACK, False, 0, 6, 1.10)
    set_style(doc.styles["Heading 1"], 16, BLUE, True, 16, 8)
    set_style(doc.styles["Heading 2"], 13, BLUE, True, 12, 6)
    set_style(doc.styles["Heading 3"], 12, DARK_BLUE, True, 8, 4)

    add_page_field(section.footer.paragraphs[0])

    add_numbering_definition(doc, "bullet", "•", 91, 91)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("ICO CHILDREN’S CODE ASSESSMENT")
    set_font(run, size=23, bold=True)

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(16)
    run = sub.add_run("Dono beta launch | Age Appropriate Design Code")
    set_font(run, size=14, color="444444")

    metadata = [
        ("Document", "Standalone assessment — separate from the Online Safety Act risk assessments and general DPIA"),
        ("Date", "6 August 2026"),
        ("Prepared for", "Dono beta launch"),
        ("Prepared by", "Midpage Legal Research"),
        ("Status", "Working paper — UK solicitor sign-off recommended before public launch"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{label}: ")
        set_font(r, bold=True)
        r = p.add_run(value)
        set_font(r)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()[7:]
    list_kind = None
    decimal_num_id = 99
    for raw in lines:
        line = raw.rstrip()
        if not line or line == "---":
            list_kind = None
            continue
        if line.startswith("# "):
            continue
        if line.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            if line[3:].startswith("3. Required changes by surface"):
                p.paragraph_format.page_break_before = True
            add_inline(p, line[3:])
            list_kind = None
            continue
        if line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline(p, line[4:])
            list_kind = None
            continue
        if re.match(r"^\d+\. ", line):
            if list_kind != "decimal":
                decimal_num_id += 1
                add_numbering_definition(
                    doc, "decimal", "%1.", decimal_num_id, decimal_num_id
                )
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.167
            apply_num(p, decimal_num_id)
            add_inline(p, re.sub(r"^\d+\. ", "", line))
            list_kind = "decimal"
            continue
        if line.startswith("- "):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.167
            apply_num(p, 91)
            add_inline(p, line[2:])
            list_kind = "bullet"
            continue
        if line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            p.paragraph_format.right_indent = Inches(0.25)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(line[2:])
            set_font(r, italic=True, color="444444")
            list_kind = None
            continue
        p = doc.add_paragraph()
        add_inline(p, line)
        list_kind = None

    props = doc.core_properties
    props.title = "Dono — ICO Children’s Code Assessment"
    props.subject = "Standalone Age Appropriate Design Code assessment for Dono"
    props.author = "Midpage Legal Research"
    props.keywords = "Dono, ICO, Children’s Code, Age Appropriate Design Code, UK GDPR"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
